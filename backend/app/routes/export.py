# Export routes for different formats
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pathlib import Path
from ..config import TEXT_DIR
from ..database import SessionLocal
from ..models import Transcript
import io
import re
import logging
from urllib.parse import quote

logger = logging.getLogger(__name__)


def sanitize_text(text: str) -> str:
    """Remove invalid XML characters (NULL bytes, control characters)"""
    if not text:
        return ""
    # Remove NULL bytes and control characters (except newline, tab, carriage return)
    # XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD]
    clean = []
    for char in text:
        code = ord(char)
        if code == 0x9 or code == 0xA or code == 0xD:
            clean.append(char)
        elif code >= 0x20 and code <= 0xD7FF:
            clean.append(char)
        elif code >= 0xE000 and code <= 0xFFFD:
            clean.append(char)
        # Skip invalid characters
    return ''.join(clean)


def get_content_disposition(filename: str, ext: str) -> str:
    """Generate proper Content-Disposition header with encoded filename"""
    # Remove original extension and add new one
    base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
    # ASCII-safe fallback filename
    ascii_name = re.sub(r'[^\w\-]', '_', base_name)
    # URL-encoded filename for UTF-8 support
    encoded_name = quote(f"{base_name}.{ext}")
    return f"attachment; filename=\"{ascii_name}.{ext}\"; filename*=UTF-8''{encoded_name}"

router = APIRouter(prefix="/export", tags=["export"])


def parse_transcript_lines(text: str):
    """Parse transcript text into segments with timestamps and text"""
    segments = []
    lines = text.strip().split('\n')
    
    # Pattern: [00:00:00 --> 00:00:05] or [00:00:00.000 --> 00:00:05.000]  Text here
    pattern = r'\[(\d{2}:\d{2}:\d{2}(?:\.\d{3})?)\s*-->\s*(\d{2}:\d{2}:\d{2}(?:\.\d{3})?)\]\s*(.+)'
    
    for line in lines:
        match = re.match(pattern, line.strip())
        if match:
            start_time = match.group(1)
            end_time = match.group(2)
            # Remove milliseconds for display
            if '.' in start_time:
                start_time = start_time.split('.')[0]
            if '.' in end_time:
                end_time = end_time.split('.')[0]
            text_content = match.group(3).strip()
            segments.append({
                'start': start_time,
                'end': end_time,
                'text': text_content
            })
        elif line.strip():
            # Line without timestamp
            segments.append({
                'start': None,
                'end': None,
                'text': line.strip()
            })
    
    return segments


def time_to_srt_format(time_str: str) -> str:
    """Convert HH:MM:SS to SRT format HH:MM:SS,000"""
    if time_str:
        return f"{time_str},000"
    return "00:00:00,000"


def time_to_seconds(time_str: str) -> float:
    """Convert HH:MM:SS to seconds"""
    if not time_str:
        return 0.0
    parts = time_str.split(':')
    hours = int(parts[0])
    minutes = int(parts[1])
    seconds = int(parts[2])
    return hours * 3600 + minutes * 60 + seconds


@router.get("/txt/{file_id}")
async def export_txt(file_id: str):
    """Export transcript as TXT file"""
    db = SessionLocal()
    try:
        transcript = db.query(Transcript).filter(Transcript.file_id == file_id).first()
        if not transcript:
            raise HTTPException(status_code=404, detail="Транскрипция не найдена")
        
        text_path = TEXT_DIR / f"{file_id}.txt"
        if not text_path.exists():
            raise HTTPException(status_code=404, detail="Файл транскрипции не найден")
        
        content = text_path.read_text(encoding='utf-8')
        filename = transcript.filename or 'transcript'
        
        buffer = io.BytesIO(content.encode('utf-8'))
        
        return StreamingResponse(
            buffer,
            media_type="application/octet-stream",
            headers={
                "Content-Disposition": get_content_disposition(filename, "txt"),
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    finally:
        db.close()


@router.get("/srt/{file_id}")
async def export_srt(file_id: str):
    """Export transcript as SRT subtitle file"""
    db = SessionLocal()
    try:
        transcript = db.query(Transcript).filter(Transcript.file_id == file_id).first()
        if not transcript:
            raise HTTPException(status_code=404, detail="Транскрипция не найдена")
        
        text_path = TEXT_DIR / f"{file_id}.txt"
        if not text_path.exists():
            raise HTTPException(status_code=404, detail="Файл транскрипции не найден")
        
        content = text_path.read_text(encoding='utf-8')
        segments = parse_transcript_lines(content)
        
        # Generate SRT format
        srt_lines = []
        for i, seg in enumerate(segments, 1):
            if seg['start'] and seg['end']:
                srt_lines.append(str(i))
                srt_lines.append(f"{time_to_srt_format(seg['start'])} --> {time_to_srt_format(seg['end'])}")
                srt_lines.append(seg['text'])
                srt_lines.append("")  # Empty line between entries
        
        srt_content = "\n".join(srt_lines)
        filename = transcript.filename or 'transcript'
        
        buffer = io.BytesIO(srt_content.encode('utf-8'))
        
        return StreamingResponse(
            buffer,
            media_type="application/octet-stream",
            headers={
                "Content-Disposition": get_content_disposition(filename, "srt"),
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    finally:
        db.close()


@router.get("/docx/{file_id}")
async def export_docx(file_id: str):
    """Export transcript as DOCX file"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx не установлен. Выполните: pip install python-docx")
    
    db = SessionLocal()
    try:
        transcript = db.query(Transcript).filter(Transcript.file_id == file_id).first()
        if not transcript:
            raise HTTPException(status_code=404, detail="Транскрипция не найдена")
        
        text_path = TEXT_DIR / f"{file_id}.txt"
        if not text_path.exists():
            raise HTTPException(status_code=404, detail="Файл транскрипции не найден")
        
        content = text_path.read_text(encoding='utf-8')
        segments = parse_transcript_lines(content)
        
        # Create DOCX document
        doc = Document()
        
        # Title - sanitize filename for XML compatibility
        title = doc.add_heading(sanitize_text(transcript.filename or 'Транскрипция'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Модель: {transcript.model or 'base'}").italic = True
        if transcript.created_at:
            meta.add_run(f" | Дата: {transcript.created_at.strftime('%d.%m.%Y %H:%M')}").italic = True
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Content
        for seg in segments:
            p = doc.add_paragraph()
            if seg['start'] and seg['end']:
                timestamp_run = p.add_run(f"[{seg['start']} --> {seg['end']}] ")
                timestamp_run.bold = True
                timestamp_run.font.size = Pt(10)
            # Sanitize text for XML compatibility
            p.add_run(sanitize_text(seg['text']))
        
        filename = transcript.filename or 'transcript'
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": get_content_disposition(filename, "docx"),
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка экспорта DOCX: {str(e)}")
    finally:
        db.close()


@router.get("/xlsx/{file_id}")
async def export_xlsx(file_id: str):
    """Export transcript as XLSX file"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl не установлен. Выполните: pip install openpyxl")
    
    db = SessionLocal()
    try:
        transcript = db.query(Transcript).filter(Transcript.file_id == file_id).first()
        if not transcript:
            raise HTTPException(status_code=404, detail="Транскрипция не найдена")
        
        text_path = TEXT_DIR / f"{file_id}.txt"
        if not text_path.exists():
            raise HTTPException(status_code=404, detail="Файл транскрипции не найден")
        
        content = text_path.read_text(encoding='utf-8')
        segments = parse_transcript_lines(content)
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Транскрипция"
        
        # Styles
        header_font = Font(bold=True, size=12, color="FFFFFF")
        header_fill = PatternFill(start_color="667eea", end_color="667eea", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Headers
        headers = ["№", "Начало", "Конец", "Длительность (сек)", "Текст"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border
        
        # Data
        for row, seg in enumerate(segments, 2):
            ws.cell(row=row, column=1, value=row-1).border = thin_border
            ws.cell(row=row, column=2, value=seg['start'] or '').border = thin_border
            ws.cell(row=row, column=3, value=seg['end'] or '').border = thin_border
            
            # Calculate duration
            if seg['start'] and seg['end']:
                duration = time_to_seconds(seg['end']) - time_to_seconds(seg['start'])
                ws.cell(row=row, column=4, value=round(duration, 1)).border = thin_border
            else:
                ws.cell(row=row, column=4, value='').border = thin_border
            
            # Sanitize text for XML compatibility
            text_cell = ws.cell(row=row, column=5, value=sanitize_text(seg['text']))
            text_cell.border = thin_border
            text_cell.alignment = Alignment(wrap_text=True)
        
        # Column widths
        ws.column_dimensions['A'].width = 6
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 80
        
        filename = transcript.filename or 'transcript'
        
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": get_content_disposition(filename, "xlsx"),
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"XLSX export error: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка экспорта XLSX: {str(e)}")
    finally:
        db.close()

